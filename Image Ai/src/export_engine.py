"""
Multi-Format Export Engine
Supports PDF, DOCX, JSON, SRT, and more
"""

from datetime import datetime
import json
from io import BytesIO
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ExportEngine:
    """Export data to multiple formats"""
    
    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def export_to_json(self, data, metadata=None):
        """
        Export data to JSON format
        
        Args:
            data: Dictionary containing extracted text, captions, translations, etc.
            metadata: Optional metadata dictionary
        
        Returns:
            JSON string
        """
        export_data = {
            "timestamp": datetime.now().isoformat(),
            "version": "1.0",
            "data": data
        }
        
        if metadata:
            export_data["metadata"] = metadata
        
        return json.dumps(export_data, indent=2, ensure_ascii=False)
    
    def export_to_srt(self, text_segments, language="en"):
        """
        Export text to SRT subtitle format
        
        Args:
            text_segments: List of text segments or single text string
            language: Language code for subtitles
        
        Returns:
            SRT formatted string
        """
        if isinstance(text_segments, str):
            # Split text into sentences for subtitles
            import re
            sentences = re.split(r'[.!?]+', text_segments)
            text_segments = [s.strip() for s in sentences if s.strip()]
        
        srt_content = []
        duration_per_segment = 3  # seconds per subtitle
        
        for idx, segment in enumerate(text_segments, 1):
            start_time = (idx - 1) * duration_per_segment
            end_time = start_time + duration_per_segment
            
            start_formatted = self._format_srt_time(start_time)
            end_formatted = self._format_srt_time(end_time)
            
            srt_content.append(f"{idx}")
            srt_content.append(f"{start_formatted} --> {end_formatted}")
            srt_content.append(segment)
            srt_content.append("")  # Blank line between subtitles
        
        return "\n".join(srt_content)
    
    def _format_srt_time(self, seconds):
        """Format seconds to SRT timestamp (HH:MM:SS,mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"
    
    def export_to_pdf(self, data, image_path=None):
        """
        Export data to PDF format with rich formatting
        
        Args:
            data: Dictionary containing text, captions, translations
            image_path: Optional path to include image in PDF
        
        Returns:
            BytesIO object containing PDF
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Container for PDF elements
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2ca02c'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph("AI Image Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Metadata table
        meta_data = [
            ['Generated:', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Version:', '1.0'],
            ['Format:', 'PDF Report']
        ]
        
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        
        story.append(meta_table)
        story.append(Spacer(1, 20))
        
        # Include image if provided
        if image_path:
            try:
                story.append(Paragraph("Source Image", heading_style))
                img = RLImage(image_path, width=4*inch, height=3*inch, kind='proportional')
                story.append(img)
                story.append(Spacer(1, 20))
            except:
                pass  # Skip if image can't be loaded
        
        # Add extracted text
        if data.get('extracted_text'):
            story.append(Paragraph("Extracted Text (OCR)", heading_style))
            text_para = Paragraph(data['extracted_text'].replace('\n', '<br/>'), styles['BodyText'])
            story.append(text_para)
            story.append(Spacer(1, 20))
        
        # Add caption
        if data.get('caption'):
            story.append(Paragraph("AI Generated Caption", heading_style))
            caption_para = Paragraph(data['caption'], styles['BodyText'])
            story.append(caption_para)
            story.append(Spacer(1, 20))
        
        # Add translation
        if data.get('translated_text') and data.get('target_language'):
            story.append(Paragraph(f"Translation ({data['target_language']})", heading_style))
            trans_para = Paragraph(data['translated_text'].replace('\n', '<br/>'), styles['BodyText'])
            story.append(trans_para)
            story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_docx(self, data, image_path=None):
        """
        Export data to DOCX format
        
        Args:
            data: Dictionary containing text, captions, translations
            image_path: Optional path to include image in document
        
        Returns:
            BytesIO object containing DOCX
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Image Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Version: 1.0")
        doc.add_paragraph()
        
        # Include image if provided
        if image_path:
            try:
                doc.add_heading('Source Image', level=1)
                doc.add_picture(image_path, width=Inches(5))
                doc.add_paragraph()
            except:
                pass  # Skip if image can't be loaded
        
        # Extracted text
        if data.get('extracted_text'):
            doc.add_heading('Extracted Text (OCR)', level=1)
            p = doc.add_paragraph(data['extracted_text'])
            p.style = 'Body Text'
            doc.add_paragraph()
        
        # Caption
        if data.get('caption'):
            doc.add_heading('AI Generated Caption', level=1)
            p = doc.add_paragraph(data['caption'])
            p.style = 'Body Text'
            doc.add_paragraph()
        
        # Translation
        if data.get('translated_text') and data.get('target_language'):
            doc.add_heading(f"Translation ({data['target_language']})", level=1)
            p = doc.add_paragraph(data['translated_text'])
            p.style = 'Body Text'
            doc.add_paragraph()
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_to_txt(self, data, include_metadata=True):
        """
        Export data to plain text format
        
        Args:
            data: Dictionary containing text, captions, translations
            include_metadata: Whether to include metadata header
        
        Returns:
            String containing formatted text
        """
        lines = []
        
        if include_metadata:
            lines.append("=" * 60)
            lines.append("AI IMAGE ANALYSIS REPORT")
            lines.append("=" * 60)
            lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"Version: 1.0")
            lines.append("=" * 60)
            lines.append("")
        
        if data.get('extracted_text'):
            lines.append("EXTRACTED TEXT (OCR)")
            lines.append("-" * 60)
            lines.append(data['extracted_text'])
            lines.append("")
        
        if data.get('caption'):
            lines.append("AI GENERATED CAPTION")
            lines.append("-" * 60)
            lines.append(data['caption'])
            lines.append("")
        
        if data.get('translated_text') and data.get('target_language'):
            lines.append(f"TRANSLATION ({data['target_language'].upper()})")
            lines.append("-" * 60)
            lines.append(data['translated_text'])
            lines.append("")
        
        return "\n".join(lines)
    
    def create_batch_report(self, batch_results):
        """
        Create a comprehensive report for batch processing
        
        Args:
            batch_results: List of dictionaries containing results from multiple images
        
        Returns:
            Dictionary with multiple export formats
        """
        report = {
            "summary": {
                "total_images": len(batch_results),
                "timestamp": datetime.now().isoformat(),
                "successful": sum(1 for r in batch_results if r.get('success')),
                "failed": sum(1 for r in batch_results if not r.get('success'))
            },
            "results": batch_results
        }
        
        return {
            "json": self.export_to_json(report),
            "txt": self._batch_to_txt(batch_results),
            "summary": report["summary"]
        }
    
    def _batch_to_txt(self, batch_results):
        """Convert batch results to text format"""
        lines = []
        lines.append("=" * 60)
        lines.append("BATCH PROCESSING REPORT")
        lines.append("=" * 60)
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Total Images: {len(batch_results)}")
        lines.append("=" * 60)
        lines.append("")
        
        for idx, result in enumerate(batch_results, 1):
            lines.append(f"IMAGE {idx}: {result.get('filename', 'Unknown')}")
            lines.append("-" * 60)
            
            if result.get('extracted_text'):
                lines.append(f"OCR Text: {result['extracted_text'][:100]}...")
            
            if result.get('caption'):
                lines.append(f"Caption: {result['caption']}")
            
            lines.append("")
        
        return "\n".join(lines)
